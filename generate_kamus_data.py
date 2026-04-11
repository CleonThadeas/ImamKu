from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import List, Tuple, Any

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title
title = doc.add_heading('Kamus Data (Data Dictionary)', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Sistem Manajemen Jadwal Imam Masjid — ImamKu')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph('')

# Helper function
def add_table_section(doc, number, table_name, description, columns):
    heading = doc.add_heading(f'{number}. Tabel {table_name}', level=2)

    desc = doc.add_paragraph(description)
    desc.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['No', 'Nama Kolom', 'Tipe Data', 'Constraint', 'Keterangan']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i] # type: ignore
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)

    for idx, col in enumerate(columns, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = col[0]
        row_cells[2].text = col[1]
        row_cells[3].text = col[2]
        row_cells[4].text = col[3]
        for cell in row_cells:
            for p in cell.paragraphs: # type: ignore
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph('')

add_table_section(doc, 1, 'users',
    'Menyimpan seluruh data akun pengguna sistem, baik Admin (Takmir Masjid) maupun Imam. Tabel ini menjadi pusat autentikasi dan otorisasi. Setiap pengguna memiliki role yang menentukan hak aksesnya — Admin dapat mengelola seluruh sistem melalui panel web, sedangkan Imam hanya dapat mengakses fitur-fitur terkait jadwal, absensi, dan pendapatan melalui web atau aplikasi mobile.',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik pengguna'),
        ('name', 'VARCHAR(255)', 'NOT NULL', 'Nama lengkap pengguna'),
        ('email', 'VARCHAR(255)', 'UNIQUE, NOT NULL', 'Alamat email untuk login'),
        ('email_verified_at', 'TIMESTAMP', 'NULLABLE', 'Waktu verifikasi email'),
        ('password', 'VARCHAR(255)', 'NOT NULL', 'Password ter-hash (bcrypt)'),
        ('role', "ENUM('admin','imam')", "DEFAULT 'imam'", 'Peran pengguna dalam sistem'),
        ('phone', 'VARCHAR(255)', 'NULLABLE', 'Nomor telepon/WhatsApp'),
        ('is_active', 'BOOLEAN', 'DEFAULT TRUE', 'Status aktif/nonaktif akun'),
        ('remember_token', 'VARCHAR(100)', 'NULLABLE', 'Token "Remember Me" Laravel'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan record'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 2, 'ramadan_seasons',
    'Menyimpan data periode/musim Ramadan. Setiap tahun Hijriyah memiliki satu record musim. Hanya satu musim yang boleh aktif (is_active = true) pada satu waktu. Musim ini menjadi induk dari seluruh jadwal, waktu sholat, dan konfigurasi fee yang berlaku pada tahun tersebut.',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik musim'),
        ('name', 'VARCHAR(255)', 'NOT NULL', 'Nama musim (contoh: "Ramadan 1447H")'),
        ('hijri_year', 'SMALLINT UNSIGNED', 'NOT NULL', 'Tahun Hijriyah'),
        ('start_date', 'DATE', 'NOT NULL', 'Tanggal mulai Ramadan'),
        ('end_date', 'DATE', 'NOT NULL', 'Tanggal berakhir Ramadan'),
        ('is_active', 'BOOLEAN', 'DEFAULT FALSE', 'Penanda musim yang sedang aktif'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 3, 'prayer_types',
    'Tabel master yang menyimpan jenis-jenis sholat yang dikelola dalam sistem. Data di tabel ini bersifat tetap (seed data) dan digunakan sebagai referensi oleh tabel schedules, prayer_times, dan fee_details. Kolom sort_order mengatur urutan tampil di antarmuka (1=Subuh hingga 6=Tarawih).',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik jenis sholat'),
        ('name', 'VARCHAR(255)', 'NOT NULL', 'Nama sholat (Subuh, Dzuhur, Ashar, Maghrib, Isya, Tarawih)'),
        ('group_code', 'CHAR(1)', 'NOT NULL', 'Kode grup untuk pengelompokan waktu'),
        ('sort_order', 'TINYINT UNSIGNED', 'NOT NULL', 'Urutan tampil di UI (1 s/d 6)'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 4, 'prayer_times',
    'Menyimpan waktu sholat harian untuk setiap musim Ramadan. Data api_time diambil secara otomatis dari API Aladhan (sinkronisasi waktu sholat berdasarkan koordinat lokasi masjid), sedangkan override_time diisi manual oleh Admin jika waktu setempat berbeda (koreksi ihtiyat). Sistem akan memprioritaskan override_time jika terisi, jika tidak maka menggunakan api_time. UNIQUE INDEX: Kombinasi (season_id, date, prayer_type_id).',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik'),
        ('season_id', 'BIGINT UNSIGNED', 'FK -> ramadan_seasons.id, CASCADE', 'Referensi ke musim Ramadan'),
        ('date', 'DATE', 'NOT NULL', 'Tanggal berlakunya waktu sholat'),
        ('prayer_type_id', 'BIGINT UNSIGNED', 'FK -> prayer_types.id, CASCADE', 'Referensi ke jenis sholat'),
        ('api_time', 'TIME', 'NULLABLE', 'Waktu sholat otomatis dari API Aladhan'),
        ('override_time', 'TIME', 'NULLABLE', 'Waktu sholat hasil koreksi manual Admin'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 5, 'schedules',
    'Tabel inti sistem yang merepresentasikan penugasan seorang Imam untuk memimpin satu jenis sholat pada tanggal tertentu. Setiap baris adalah satu "slot" jadwal. Admin mengisi kolom user_id saat menugaskan Imam. Kolom ini bisa bernilai NULL jika slot belum terisi. Tabel ini berelasi langsung dengan tabel attendances (absensi) dan swap_requests (tukar jadwal). UNIQUE INDEX: (season_id, date, prayer_type_id). INDEX: (date, user_id).',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik jadwal'),
        ('season_id', 'BIGINT UNSIGNED', 'FK -> ramadan_seasons.id, CASCADE', 'Referensi ke musim'),
        ('date', 'DATE', 'NOT NULL', 'Tanggal penugasan'),
        ('prayer_type_id', 'BIGINT UNSIGNED', 'FK -> prayer_types.id, CASCADE', 'Jenis sholat yang ditugaskan'),
        ('user_id', 'BIGINT UNSIGNED', 'FK -> users.id, NULLABLE', 'Imam yang ditugaskan (NULL jika belum diisi)'),
        ('notes', 'TEXT', 'NULLABLE', 'Catatan tambahan dari Admin'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 6, 'attendances',
    'Menyimpan bukti kehadiran (absensi) Imam pada jadwal sholat yang ditugaskan. Saat Imam melakukan Check-In melalui aplikasi mobile, ia mengunggah foto bukti kehadiran (proof_path). Status awal adalah pending, kemudian Admin melakukan validasi untuk mengubahnya menjadi approved (disetujui) atau rejected (ditolak). Status approved menjadi syarat agar fee/insentif Imam dicairkan.',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik absensi'),
        ('schedule_id', 'BIGINT UNSIGNED', 'FK -> schedules.id, CASCADE', 'Jadwal yang diabsensi'),
        ('proof_path', 'VARCHAR(255)', 'NULLABLE', 'Path file foto bukti kehadiran di storage'),
        ('status', "ENUM('pending','approved','rejected')", "DEFAULT 'pending'", 'Status validasi oleh Admin'),
        ('notes', 'TEXT', 'NULLABLE', 'Catatan dari Admin saat mereview bukti'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu Imam melakukan check-in'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 7, 'swap_requests',
    'Mencatat seluruh permintaan tukar jadwal antar Imam. Alurnya: Imam A mengajukan permintaan tukar untuk schedule_id (jadwal miliknya). Imam B yang bersedia akan menawarkan target_schedule_id (jadwal milik Imam B sebagai gantinya). Jika disetujui (accepted), sistem otomatis membalikkan user_id di kedua jadwal. Status expired dikenakan jika permintaan melewati batas waktu toleransi (H-2 jam sebelum waktu sholat).',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik permintaan swap'),
        ('schedule_id', 'BIGINT UNSIGNED', 'FK -> schedules.id, CASCADE', 'Jadwal asal yang diminta tukar'),
        ('target_schedule_id', 'BIGINT UNSIGNED', 'FK -> schedules.id, NULLABLE, CASCADE', 'Jadwal pengganti yang ditawarkan'),
        ('requester_id', 'BIGINT UNSIGNED', 'FK -> users.id, CASCADE', 'Imam yang mengajukan permintaan'),
        ('status', "ENUM('pending','accepted','rejected','expired')", "DEFAULT 'pending'", 'Status permintaan swap'),
        ('processed_at', 'DATETIME', 'NULLABLE', 'Waktu swap diproses/disetujui/ditolak'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pengajuan'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 8, 'fee_configs',
    'Menyimpan konfigurasi utama sistem insentif (fee) per musim Ramadan. Admin dapat mengatur mode perhitungan fee (per_schedule = dihitung per sholat, per_day = dihitung per hari). Kolom is_auto_approve_attendance jika diaktifkan akan otomatis meng-approve absensi Imam tanpa perlu validasi manual Admin, sehingga fee langsung terhitung.',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik konfigurasi'),
        ('season_id', 'BIGINT UNSIGNED', 'FK -> ramadan_seasons.id, CASCADE', 'Musim yang berlaku'),
        ('mode', "ENUM('per_schedule','per_day')", "DEFAULT 'per_schedule'", 'Mode perhitungan fee'),
        ('is_enabled', 'BOOLEAN', 'DEFAULT FALSE', 'Apakah fitur fee diaktifkan'),
        ('is_auto_approve_attendance', 'BOOLEAN', 'DEFAULT FALSE', 'Otomatis approve absensi tanpa review'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 9, 'fee_details',
    'Menyimpan rincian nominal tarif per jenis sholat. Tabel ini merupakan anak dari fee_configs. Setiap jenis sholat dapat memiliki nominal yang berbeda (misal: Tarawih lebih besar dari Subuh). Nilai amount digunakan oleh FeeService untuk menghitung total pendapatan Imam secara on-the-fly berdasarkan jumlah absensi yang di-approve.',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik rincian tarif'),
        ('fee_config_id', 'BIGINT UNSIGNED', 'FK -> fee_configs.id, CASCADE', 'Referensi ke induk konfigurasi'),
        ('prayer_type_id', 'BIGINT UNSIGNED', 'FK -> prayer_types.id, NULLABLE', 'Jenis sholat yang diberi tarif'),
        ('amount', 'DECIMAL(12,2)', 'DEFAULT 0', 'Nominal fee dalam Rupiah'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 10, 'notification_logs',
    'Mencatat seluruh riwayat notifikasi yang dikirim oleh sistem kepada setiap Imam. Setiap kali sistem mengirimkan reminder sholat, pemberitahuan swap, atau broadcast dari Admin, satu record dibuat di tabel ini. Kolom channel mencatat kanal pengiriman (database, whatsapp, email), sedangkan payload berisi data JSON lengkap isi notifikasi. Tabel ini juga berfungsi sebagai audit trail untuk memastikan tidak ada notifikasi yang terlewat. INDEX: (user_id, schedule_id, type).',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik log'),
        ('user_id', 'BIGINT UNSIGNED', 'FK -> users.id, CASCADE', 'Imam penerima notifikasi'),
        ('schedule_id', 'BIGINT UNSIGNED', 'FK -> schedules.id, NULLABLE', 'Jadwal terkait (jika relevan)'),
        ('channel', 'VARCHAR(255)', 'NOT NULL', 'Kanal pengiriman (database, whatsapp, email)'),
        ('type', 'VARCHAR(255)', 'NOT NULL', 'Tipe notifikasi (reminder, swap_request, broadcast)'),
        ('payload', 'JSON', 'NULLABLE', 'Data lengkap isi notifikasi dalam format JSON'),
        ('status', 'VARCHAR(50)', "DEFAULT 'pending'", 'Status pengiriman (sent, failed, pending)'),
        ('error_message', 'TEXT', 'NULLABLE', 'Pesan error jika gagal terkirim'),
        ('sent_at', 'DATETIME', 'NULLABLE', 'Waktu notifikasi berhasil terkirim'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu record dibuat'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 11, 'notification_configs',
    'Tabel konfigurasi global (singleton, hanya memiliki 1 baris data) untuk mengatur perilaku sistem notifikasi. Admin dapat menentukan berapa menit sebelum waktu sholat sistem akan mengirim Reminder Pertama (reminder_1_minutes, default 90 menit) dan apakah Reminder Kedua diaktifkan (enable_reminder_2, default mati). Kolom channels mengatur kanal mana saja yang aktif.',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik'),
        ('channels', 'VARCHAR(255)', "DEFAULT 'database,whatsapp'", 'Kanal pengiriman aktif (dipisah koma)'),
        ('reminder_1_minutes', 'INT', 'DEFAULT 90', 'Menit sebelum sholat untuk Reminder Pertama'),
        ('enable_reminder_2', 'BOOLEAN', 'DEFAULT FALSE', 'Aktifkan pengingat kedua (ya/tidak)'),
        ('reminder_2_minutes', 'INT', 'DEFAULT 30', 'Menit sebelum sholat untuk Reminder Kedua'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

add_table_section(doc, 12, 'personal_access_tokens',
    'Tabel bawaan Laravel Sanctum yang menyimpan token autentikasi API untuk aplikasi mobile Flutter. Setiap kali Imam berhasil login melalui aplikasi mobile, sistem membuat satu token baru di tabel ini. Token ini dikirimkan sebagai header Bearer di setiap request API berikutnya untuk memverifikasi identitas Imam. Token dapat memiliki masa kadaluarsa (expires_at) dan mencatat waktu terakhir digunakan (last_used_at).',
    [
        ('id', 'BIGINT UNSIGNED', 'PK, AUTO_INCREMENT', 'ID unik token'),
        ('tokenable_type', 'VARCHAR(255)', 'NOT NULL', 'Tipe model pemilik (App\\Models\\User)'),
        ('tokenable_id', 'BIGINT UNSIGNED', 'NOT NULL', 'ID user pemilik token'),
        ('name', 'VARCHAR(255)', 'NOT NULL', 'Nama label token (default: "mobile")'),
        ('token', 'VARCHAR(64)', 'UNIQUE', 'Hash SHA-256 dari plain-text token'),
        ('abilities', 'TEXT', 'NULLABLE', 'Daftar izin/kemampuan token'),
        ('last_used_at', 'TIMESTAMP', 'NULLABLE', 'Waktu terakhir token digunakan'),
        ('expires_at', 'TIMESTAMP', 'NULLABLE', 'Waktu kadaluarsa token'),
        ('created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan'),
        ('updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu perubahan terakhir'),
    ])

output_path = r'd:\laragon\www\ImamKuRamadanSchedule\Kamus_Data_ImamKu.docx'
doc.save(output_path)
print(f'File saved to: {output_path}')
